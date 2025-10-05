# Copyright (c) 2025, Coale Tech and contributors
# For license information, please see license.txt

import frappe
from frappe import _
from frappe.model.document import Document
from frappe.utils import nowdate, getdate
import re
from docx import Document as DocxDocument
from docx.shared import Inches
import io

class LegalDocumentTemplate(Document):
	def validate(self):
		"""Validate the document template"""
		if self.template_content:
			# Check for merge field syntax
			merge_fields = re.findall(r'\{([^}]+)\}', self.template_content)
			if merge_fields:
				valid_fields = [f.field_name for f in self.merge_fields]
				invalid_fields = [field for field in merge_fields if field not in valid_fields]
				if invalid_fields:
					frappe.msgprint(_("The following merge fields are used in content but not defined: {0}").format(", ".join(invalid_fields)),
						alert=True)

	def on_update(self):
		"""Clear cache when template is updated"""
		frappe.clear_cache()

@frappe.whitelist()
def generate_document(template_name, reference_doc=None, merge_data=None):
	"""Generate document from template"""
	template = frappe.get_doc("Legal Document Template", template_name)
	if not template:
		frappe.throw(_("Template not found"))

	# Get merge data
	if not merge_data:
		merge_data = get_merge_data(template, reference_doc)

	# Process template content
	content = template.template_content
	for field in template.merge_fields:
		field_value = merge_data.get(field.field_name, field.default_value or "")
		if field.field_type == "Currency" and field_value:
			field_value = frappe.utils.fmt_money(field_value)
		elif field.field_type == "Date" and field_value:
			field_value = getdate(field_value).strftime('%d/%m/%Y')

		content = content.replace(f"{{{field.field_name}}}", str(field_value))

	# Generate output based on format
	if template.output_format == "DOCX":
		return generate_docx(content, template.template_name)
	elif template.output_format == "PDF":
		return generate_pdf(content, template.template_name)
	else:
		return content

def get_merge_data(template, reference_doc):
	"""Get merge data from reference document"""
	merge_data = {}

	if not reference_doc:
		return merge_data

	# Get data from reference document
	doc = frappe.get_doc(reference_doc)
	doc_dict = doc.as_dict()

	for field in template.merge_fields:
		if field.field_source == "Manual Entry":
			continue
		elif field.field_source in ["Lead", "Deal", "Organization", "Contact"]:
			# Map source to actual field
			source_mapping = {
				"Lead": doc_dict,
				"Deal": doc_dict,
				"Organization": doc_dict.get("organization_doc", {}),
				"Contact": doc_dict.get("contact_doc", {})
			}
			source_data = source_mapping.get(field.field_source, {})
			merge_data[field.field_name] = source_data.get(field.field_name, "")
		else:
			merge_data[field.field_name] = doc_dict.get(field.field_name, "")

	# Add common fields
	merge_data.update({
		"current_date": nowdate().strftime('%d/%m/%Y'),
		"company_name": frappe.db.get_single_value("Company", "company_name") or "",
		"law_firm_address": frappe.db.get_single_value("Company", "address") or "",
	})

	return merge_data

def generate_docx(content, title):
	"""Generate DOCX document"""
	doc = DocxDocument()
	doc.add_heading(title, 0)

	# Split content by paragraphs and add to document
	paragraphs = content.split('\n\n')
	for para in paragraphs:
		if para.strip():
			doc.add_paragraph(para.strip())

	# Save to bytes
	docx_bytes = io.BytesIO()
	doc.save(docx_bytes)
	docx_bytes.seek(0)

	return docx_bytes.getvalue()

def generate_pdf(content, title):
	"""Generate PDF document using WeasyPrint"""
	try:
		from weasyprint import HTML, CSS
	except ImportError:
		frappe.throw(_("WeasyPrint is required for PDF generation"))

	html_content = f"""
	<html>
	<head>
		<title>{title}</title>
		<style>
			body {{ font-family: Arial, sans-serif; margin: 40px; }}
			h1 {{ color: #333; }}
			p {{ line-height: 1.6; }}
		</style>
	</head>
	<body>
		<h1>{title}</h1>
		{content.replace(chr(10), '<br>')}
	</body>
	</html>
	"""

	html_doc = HTML(string=html_content)
	pdf_bytes = html_doc.write_pdf()

	return pdf_bytes

@frappe.whitelist()
def get_templates_by_type(document_type, practice_area=None):
	"""Get templates by document type and practice area"""
	filters = {"document_type": document_type, "is_active": 1}
	if practice_area:
		filters["practice_area"] = practice_area

	templates = frappe.get_all("Legal Document Template",
		filters=filters,
		fields=["name", "template_name", "description"]
	)

	return templates

@frappe.whitelist()
def auto_generate_documents():
	"""Auto-generate documents based on triggers"""
	templates = frappe.get_all("Legal Document Template",
		filters={"auto_generate": 1, "is_active": 1},
		fields=["name", "trigger_event"]
	)

	for template in templates:
		if template.trigger_event == "Lead Created":
			generate_for_new_leads(template.name)
		elif template.trigger_event == "Deal Won":
			generate_for_won_deals(template.name)
		elif template.trigger_event == "Case Created":
			generate_for_new_cases(template.name)

def generate_for_new_leads(template_name):
	"""Generate documents for newly created leads"""
	# Get leads created in the last hour
	leads = frappe.db.sql("""
		SELECT name
		FROM `tabLegal CRM Lead`
		WHERE creation >= DATE_SUB(NOW(), INTERVAL 1 HOUR)
		AND status = 'New'
	""", as_dict=True)

	for lead in leads:
		try:
			document = generate_document(template_name, {"doctype": "Legal CRM Lead", "name": lead.name})
			# Save generated document as attachment
			save_generated_document(document, template_name, lead.name, "Legal CRM Lead")
		except Exception as e:
			frappe.log_error(f"Auto-generation failed for lead {lead.name}: {str(e)}")

def generate_for_won_deals(template_name):
	"""Generate documents for won deals"""
	deals = frappe.db.sql("""
		SELECT name
		FROM `tabLegal CRM Deal`
		WHERE status = 'Won'
		AND closed_date >= DATE_SUB(NOW(), INTERVAL 1 HOUR)
	""", as_dict=True)

	for deal in deals:
		try:
			document = generate_document(template_name, {"doctype": "Legal CRM Deal", "name": deal.name})
			save_generated_document(document, template_name, deal.name, "Legal CRM Deal")
		except Exception as e:
			frappe.log_error(f"Auto-generation failed for deal {deal.name}: {str(e)}")

def generate_for_new_cases(template_name):
	"""Generate documents for new cases"""
	cases = frappe.db.sql("""
		SELECT name
		FROM `tabLegal Case`
		WHERE creation >= DATE_SUB(NOW(), INTERVAL 1 HOUR)
	""", as_dict=True)

	for case in cases:
		try:
			document = generate_document(template_name, {"doctype": "Legal Case", "name": case.name})
			save_generated_document(document, template_name, case.name, "Legal Case")
		except Exception as e:
			frappe.log_error(f"Auto-generation failed for case {case.name}: {str(e)}")

def save_generated_document(document_content, template_name, reference_name, reference_doctype):
	"""Save generated document as attachment"""
	template = frappe.get_doc("Legal Document Template", template_name)

	file_name = f"{template.document_type}_{reference_name}_{nowdate()}.{template.output_format.lower()}"

	# Create file document
	file_doc = frappe.new_doc("File")
	file_doc.file_name = file_name
	file_doc.attached_to_doctype = reference_doctype
	file_doc.attached_to_name = reference_name
	file_doc.content = document_content
	file_doc.insert(ignore_permissions=True)

	return file_doc.name

# Default templates
DEFAULT_ENGAGEMENT_LETTER = """
[Law Firm Letterhead]

{current_date}

{client_name}
{client_address}

Dear {client_salutation} {client_first_name} {client_last_name},

RE: ENGAGEMENT LETTER - {practice_area}

We are pleased to confirm our appointment as your legal counsel in relation to {case_description}.

1. SCOPE OF SERVICES
We shall provide legal services in the following areas:
- {practice_area} matters
- Legal advice and representation
- Court appearances and filings

2. FEES AND BILLING
Our professional fees shall be charged at the rate of KES {hourly_rate} per hour.
Disbursements and expenses shall be charged separately.

3. TRUST ACCOUNT
Client funds will be held in our trust account in accordance with the Advocates Act.

4. CONFIDENTIALITY
All communications and information shared will be treated as confidential.

Please sign and return a copy of this letter to confirm your acceptance of our engagement.

Yours faithfully,

{Lawyer Name}
{Lawyer Position}
{Company Name}
"""

def create_default_templates():
	"""Create default document templates"""
	if not frappe.db.exists("Legal Document Template", "Engagement Letter - Default"):
		template = frappe.new_doc("Legal Document Template")
		template.template_name = "Engagement Letter - Default"
		template.document_type = "Engagement Letter"
		template.description = "Standard engagement letter template"
		template.template_content = DEFAULT_ENGAGEMENT_LETTER
		template.output_format = "DOCX"
		template.generate_pdf = True
		template.auto_generate = True
		template.trigger_event = "Deal Won"
		template.is_active = 1

		# Add merge fields
		merge_fields = [
			{"field_name": "client_name", "field_label": "Client Name", "field_type": "Text", "field_source": "Lead", "is_required": True},
			{"field_name": "client_address", "field_label": "Client Address", "field_type": "Text", "field_source": "Lead"},
			{"field_name": "client_salutation", "field_label": "Client Salutation", "field_type": "Text", "field_source": "Lead"},
			{"field_name": "client_first_name", "field_label": "Client First Name", "field_type": "Text", "field_source": "Lead"},
			{"field_name": "client_last_name", "field_label": "Client Last Name", "field_type": "Text", "field_source": "Lead"},
			{"field_name": "practice_area", "field_label": "Practice Area", "field_type": "Text", "field_source": "Lead"},
			{"field_name": "case_description", "field_label": "Case Description", "field_type": "Text", "field_source": "Manual Entry"},
			{"field_name": "hourly_rate", "field_label": "Hourly Rate", "field_type": "Currency", "field_source": "Manual Entry", "default_value": "5000"},
			{"field_name": "lawyer_name", "field_label": "Lawyer Name", "field_type": "Text", "field_source": "Manual Entry"},
			{"field_name": "lawyer_position", "field_label": "Lawyer Position", "field_type": "Text", "field_source": "Manual Entry"},
			{"field_name": "company_name", "field_label": "Company Name", "field_type": "Text", "field_source": "Manual Entry"},
		]

		for field_data in merge_fields:
			field = template.append("merge_fields", field_data)
			field.field_name = field_data["field_name"]
			field.field_label = field_data["field_label"]
			field.field_type = field_data["field_type"]
			field.field_source = field_data["field_source"]
			field.is_required = field_data.get("is_required", False)
			field.default_value = field_data.get("default_value", "")

		template.insert(ignore_permissions=True)

	frappe.db.commit()


# Standalone function for testing
def process_merge_fields(template_content, reference_data):
	"""Process merge fields in template content for testing"""
	if not template_content or not reference_data:
		return template_content

	content = template_content
	for key, value in reference_data.items():
		content = content.replace(f"{{{key}}}", str(value))

	return content